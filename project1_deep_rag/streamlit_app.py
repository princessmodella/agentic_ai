# project1_deep_rag/app.py
import os
import io
import streamlit as st
from streamlit_option_menu import option_menu
from indexer import add_to_collection, retrieve, reset_collection
import fitz  # PyMuPDF for PDF
import docx
from reportlab.pdfgen import canvas
from docx import Document

st.set_page_config(page_title="Deep RAG Engine", page_icon="🤖", layout="wide")

# ---------------- Utility: Download Buttons ---------------- #
def download_buttons(content, filename_prefix="report"):
    """Generate side-by-side download buttons for TXT, PDF, DOCX + copy option."""
    col1, col2, col3, col4 = st.columns(4)

    # TXT
    txt_data = io.BytesIO(content.encode("utf-8"))
    with col1:
        st.download_button(
            "⬇️ TXT",
            data=txt_data,
            file_name=f"{filename_prefix}.txt",
            mime="text/plain"
        )

    # PDF
    pdf_data = io.BytesIO()
    c = canvas.Canvas(pdf_data)
    text_object = c.beginText(40, 800)
    for line in content.split("\n"):
        text_object.textLine(line)
    c.drawText(text_object)
    c.save()
    pdf_data.seek(0)
    with col2:
        st.download_button(
            "⬇️ PDF",
            data=pdf_data,
            file_name=f"{filename_prefix}.pdf",
            mime="application/pdf"
        )

    # DOCX
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    docx_data = io.BytesIO()
    doc.save(docx_data)
    docx_data.seek(0)
    with col3:
        st.download_button(
            "⬇️ DOCX",
            data=docx_data,
            file_name=f"{filename_prefix}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Copy to Clipboard
    with col4:
        st.text_area("📋 Copy", content, height=100)


# ---------------- Sidebar Navigation ---------------- #
with st.sidebar:
    selected = option_menu(
        "Deep RAG Engine",
        ["Home", "Ask Question", "Upload & Summarize"],
        icons=["house", "message-circle", "file-text"],
        menu_icon="robot"
    )

# ---------------- Home Page ---------------- #
if selected == "Home":
    st.title("🤖 Deep RAG Engine")
    st.image("project1_deep_rag/OIP (1).webp", width='stretch')
    st.write(
        "AI-powered tool to answer questions from web, URLs, or text content. "
        "You can also upload **multiple documents** for summarization and download reports."
    )

# ---------------- Ask Question ---------------- #
elif selected == "Ask Question":
    st.title("💬 Ask a Question")
    question = st.text_input("Type your question here")

    if st.button("Get Answer") and question.strip():
        with st.spinner("Searching..."):
            context = retrieve(question)

        st.subheader("📝 Answer")
        # Ensure context is a string
        if isinstance(context, list):
            context = " ".join(context)

        bullets = context.split(". ")
        formatted_answer = "\n".join([f"- {b.strip()}" for b in bullets if b.strip()])
        st.write(formatted_answer)
        st.write("📥 **Download Answer**")
        download_buttons(formatted_answer, "answer")

# ---------------- Upload & Summarize ---------------- #
elif selected == "Upload & Summarize":
    st.title("📂 Upload & Summarize")
    uploaded_files = st.file_uploader(
        "Upload one or more PDF, TXT, or DOCX files",
        type=["pdf", "txt", "docx"],
        accept_multiple_files=True
    )

    if uploaded_files:
        combined_text = ""
        for uploaded_file in uploaded_files:
            text = ""
            if uploaded_file.type == "application/pdf":
                pdf_doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                for page in pdf_doc:
                    text += page.get_text()
            elif uploaded_file.type == "text/plain":
                text = uploaded_file.read().decode("utf-8")
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = docx.Document(uploaded_file)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            combined_text += text + "\n\n"

        if combined_text.strip():
            add_to_collection(combined_text)
            st.subheader("📄 Combined Document Summary")
            summary_context = retrieve("Summarize this document collection", top_k=5)
            if isinstance(summary_context, list):
                summary_context = " ".join(summary_context)
            bullets = summary_context.split(". ")
            formatted_summary = "\n".join([f"- {b.strip()}" for b in bullets if b.strip()])
            st.write(formatted_summary)
            st.write("📥 **Download Summary**")
            download_buttons(formatted_summary, "summary")
        else:
            st.error("❌ No readable text found in uploaded files.")
